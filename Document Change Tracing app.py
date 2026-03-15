import streamlit as st
import pandas as pd
import docx
import re
import difflib
from openpyxl import Workbook
from openpyxl.styles import Font
from io import BytesIO

st.set_page_config(page_title="Clause Change Log Tool", layout="wide")
st.title("📄 Clause Change Log Generator – Full Clause + Highlights")

# -----------------------------
# Extract clauses including multi-line content
# -----------------------------
def extract_clauses(file):
    doc = docx.Document(file)
    clauses = {}
    pattern = r'(\d+\.\d+(\.\d+)*)'
    current_clause = None
    current_text = []

    # Combine paragraphs and tables
    for block in doc.paragraphs + [cell for table in doc.tables for row in table.rows for cell in row.cells]:
        text = block.text.strip()
        if not text:
            continue
        match = re.match(pattern, text)
        if match:
            # Save previous clause
            if current_clause:
                clauses[current_clause] = "\n".join(current_text).strip()
            # Start new clause
            current_clause = match.group(1)
            current_text = [text]
        else:
            if current_clause:
                current_text.append(text)
    # Save last clause
    if current_clause:
        clauses[current_clause] = "\n".join(current_text).strip()
    return clauses

# -----------------------------
# Word-level diff for modified clauses
# -----------------------------
def build_diff(before, after):
    # Remove clause number prefix for clarity
    before_clean = re.sub(r'^\d+(\.\d+)*\s*', '', before)
    after_clean = re.sub(r'^\d+(\.\d+)*\s*', '', after)

    before_words = before_clean.split()
    after_words = after_clean.split()
    diff = list(difflib.ndiff(before_words, after_words))
    result = []

    for d in diff:
        if d.startswith("+ "):
            result.append(("added", d[2:]))
        elif d.startswith("- "):
            result.append(("removed", d[2:]))
        elif d.startswith("  "):
            result.append(("same", d[2:]))
    return result

# -----------------------------
# File Upload
# -----------------------------
before_file = st.file_uploader("Upload BEFORE Document", type=["docx"])
after_file = st.file_uploader("Upload AFTER Document", type=["docx"])

if before_file and after_file:
    before_clauses = extract_clauses(before_file)
    after_clauses = extract_clauses(after_file)
    document_name = after_file.name
    all_clauses = set(before_clauses.keys()).union(set(after_clauses.keys()))
    results = []

    for clause in sorted(all_clauses):
        before_text = before_clauses.get(clause, "")
        after_text = after_clauses.get(clause, "")

        if clause in before_clauses and clause not in after_clauses:
            status = "Removed"
            rtype = "removed"
            rcontent = before_text

        elif clause not in before_clauses and clause in after_clauses:
            status = "New Clause Added"
            rtype = "added"
            rcontent = after_text

        elif before_text != after_text:
            status = "Statement Modified / Revised"
            rtype = "modified"
            rcontent = build_diff(before_text, after_text)
        else:
            continue

        results.append([
            document_name,
            before_text,
            after_text,
            status,
            rtype,
            rcontent
        ])

    # -----------------------------
    # Build DataFrame with SR #
    # -----------------------------
    df = pd.DataFrame(results, columns=[
        "Document Name",
        "Before Clause",
        "After Clause",
        "Status",
        "Remark Type",
        "Remarks Content"
    ])
    df.reset_index(inplace=True)
    df.rename(columns={'index':'SR #'}, inplace=True)
    df['SR #'] = df['SR #'] + 1

    # Format Remarks for Streamlit preview
    def format_remarks(row):
        rtype = row['Remark Type']
        content = row['Remarks Content']
        if rtype in ["removed", "added"]:
            return content
        elif rtype == "modified":
            # show added/removed words inline
            return " ".join([w if t=="same" else w for t,w in content])
        return content

    df['Remarks'] = df.apply(format_remarks, axis=1)
    st.subheader("📊 Change Log Preview")
    st.dataframe(df[['SR #','Document Name','Before Clause','After Clause','Status','Remarks']], use_container_width=True)

    # -----------------------------
    # Create Excel with color
    # -----------------------------
    wb = Workbook()
    ws = wb.active
    ws.title = "Change Log"
    ws.append(['SR #','Document Name','Before Clause','After Clause','Status','Remarks'])

    red = Font(color="FF0000", strike=True)
    green = Font(color="008000")
    blue = Font(color="0000FF")

    for idx, row in df.iterrows():
        sr, docname, before_c, after_c, status, rtype, rcontent, remarks = row
        ws.append([sr, docname, before_c, after_c, status, ""])
        remarks_cell = ws.cell(row=ws.max_row, column=6)

        if rtype == "removed":
            remarks_cell.value = rcontent
            remarks_cell.font = red
        elif rtype == "added":
            remarks_cell.value = rcontent
            remarks_cell.font = green
        elif rtype == "modified":
            remarks_cell.value = " ".join([w for t,w in rcontent])
            remarks_cell.font = blue

    # -----------------------------
    # Download Excel
    # -----------------------------
    buffer = BytesIO()
    wb.save(buffer)
    st.download_button(
        label="📥 Download Excel Change Log",
        data=buffer.getvalue(),
        file_name="Clause_Change_Log.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
